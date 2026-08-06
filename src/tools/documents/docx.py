"""
DOCX Parser
Part of SOVEREIGN PYTHON LLM ENGINE

Parse Microsoft Word documents (.docx).
"""

from typing import Any
from pathlib import Path
from dataclasses import dataclass
import io

from ...core.evidence import WORMLedger


@dataclass
class DOCXDocument:
    """Parsed DOCX document"""
    text: str
    paragraphs: list[str]
    tables: list[list[list[str]]]
    metadata: dict[str, Any]


class DOCXParser:
    """
    DOCX document parser using python-docx.
    """

    def __init__(self, worm_ledger: WORMLedger | None = None):
        """
        Initialize DOCX parser.

        Args:
            worm_ledger: Optional WORM ledger
        """
        self.worm_ledger = worm_ledger

    async def parse(self, file_path: Path) -> DOCXDocument:
        """
        Parse DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            DOCXDocument with extracted content
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required. Install: pip install python-docx")

        import asyncio

        def _sync_parse():
            doc = Document(file_path)

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract text
            text = "\n\n".join(paragraphs)

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            # Extract metadata
            metadata = {}
            if doc.core_properties:
                props = doc.core_properties
                metadata = {
                    "title": props.title or "",
                    "author": props.author or "",
                    "subject": props.subject or "",
                    "keywords": props.keywords or "",
                    "created": props.created.isoformat() if props.created else None,
                    "modified": props.modified.isoformat() if props.modified else None,
                }

            return DOCXDocument(
                text=text,
                paragraphs=paragraphs,
                tables=tables,
                metadata=metadata
            )

        return await asyncio.to_thread(_sync_parse)

    async def parse_bytes(self, docx_bytes: bytes) -> DOCXDocument:
        """
        Parse DOCX from bytes.

        Args:
            docx_bytes: DOCX file bytes

        Returns:
            DOCXDocument
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required")

        import asyncio

        def _sync_parse():
            doc = Document(io.BytesIO(docx_bytes))

            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            text = "\n\n".join(paragraphs)

            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            metadata = {}
            if doc.core_properties:
                props = doc.core_properties
                metadata = {
                    "title": props.title or "",
                    "author": props.author or "",
                }

            return DOCXDocument(
                text=text,
                paragraphs=paragraphs,
                tables=tables,
                metadata=metadata
            )

        return await asyncio.to_thread(_sync_parse)


# Tool registration helper
async def parse_docx_tool(file_path: str) -> dict:
    """
    Tool wrapper for DOCX parsing.

    Args:
        file_path: Path to DOCX file

    Returns:
        Dictionary with parsed content
    """
    parser = DOCXParser()
    result = await parser.parse(Path(file_path))

    return {
        "text": result.text,
        "paragraphs_count": len(result.paragraphs),
        "tables_count": len(result.tables),
        "metadata": result.metadata
    }
